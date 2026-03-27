#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成简历Word文档 - 保持网页模板布局
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

def create_resume():
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
    
    # 创建表格布局（左侧边栏 + 右侧内容）
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    
    # 设置列宽 - 左侧280px约等于2.9英寸，右侧620px约等于6.4英寸
    table.columns[0].width = Inches(2.9)
    table.columns[1].width = Inches(6.1)
    
    # 获取单元格
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    
    # 设置行高（让表格占满页面）
    table.rows[0].height = Inches(11)
    
    # ========== 左侧边栏（薄荷绿背景 #39c5bb）==========
    # 设置左侧单元格背景色
    shading_elm = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" fill="39c5bb"/>')
    left_cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # 清空默认段落
    left_cell.text = ""
    
    # 头像区域
    p = left_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(0)
    run = p.add_run("宋")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(57, 197, 187)
    
    # 姓名
    p = left_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(10)
    p.space_after = Pt(0)
    run = p.add_run("宋浩然")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 职位
    p = left_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(5)
    p.space_after = Pt(20)
    run = p.add_run("游戏客户端开发工程师\n（UE，U3D）")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 个人简介标题
    p = left_cell.add_paragraph()
    p.space_before = Pt(15)
    p.space_after = Pt(8)
    run = p.add_run("个人简介")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 个人简介内容
    p = left_cell.add_paragraph()
    p.space_after = Pt(15)
    run = p.add_run("3年游戏客户端开发经验，精通Unity和UE4/5双引擎，专注Gameplay系统架构与性能优化。曾主导PVP RTS游戏核心系统开发，具备全栈开发能力（客户端+服务器）。热爱技术钻研，持续关注图形渲染与AI行为系统。")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 联系方式标题
    p = left_cell.add_paragraph()
    p.space_before = Pt(15)
    p.space_after = Pt(8)
    run = p.add_run("联系方式")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 联系方式内容
    contact_info = [
        ("邮箱", "kirino8214@gmail.com"),
        ("电话", "19539483901"),
        ("个人网站", "pleiadorum.cn"),
    ]
    for label, value in contact_info:
        p = left_cell.add_paragraph()
        p.space_after = Pt(5)
        run = p.add_run(f"{label}：")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 教育背景标题
    p = left_cell.add_paragraph()
    p.space_before = Pt(20)
    p.space_after = Pt(8)
    run = p.add_run("教育背景")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 教育背景内容
    p = left_cell.add_paragraph()
    p.space_after = Pt(5)
    run = p.add_run("湖南大学")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    p = left_cell.add_paragraph()
    p.space_after = Pt(3)
    run = p.add_run("数字媒体技术 - 本科")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    p = left_cell.add_paragraph()
    run = p.add_run("2019.06 - 2023.07")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 技能标签标题
    p = left_cell.add_paragraph()
    p.space_before = Pt(20)
    p.space_after = Pt(10)
    run = p.add_run("技能标签")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 技能标签内容 - 分两行显示
    skills_line1 = ["Unity", "UE4", "UE5", "C#", "C++", "Lua"]
    skills_line2 = ["Python", "Gameplay", "性能优化", "帧同步", "Shader", "Git"]
    
    p = left_cell.add_paragraph()
    p.space_after = Pt(5)
    for i, skill in enumerate(skills_line1):
        run = p.add_run(skill)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        if i < len(skills_line1) - 1:
            run = p.add_run("  •  ")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
    
    p = left_cell.add_paragraph()
    for i, skill in enumerate(skills_line2):
        run = p.add_run(skill)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        if i < len(skills_line2) - 1:
            run = p.add_run("  •  ")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
    
    # ========== 右侧内容区域（白色背景）==========
    right_cell.text = ""
    
    # 工作经历标题
    p = right_cell.add_paragraph()
    p.space_after = Pt(12)
    run = p.add_run("工作经历")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(57, 197, 187)
    
    # 添加下划线效果（用段落边框模拟）
    p = right_cell.add_paragraph()
    p.space_after = Pt(15)
    run = p.add_run("─" * 50)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(57, 197, 187)
    
    # 工作经历1
    p = right_cell.add_paragraph()
    p.space_after = Pt(3)
    run = p.add_run("2025.02 - 至今")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    p = right_cell.add_paragraph()
    p.space_after = Pt(8)
    run = p.add_run("UE开发")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 51, 51)
    run = p.add_run("  |  ")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(153, 153, 153)
    run = p.add_run("上海冒险岛网络")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(57, 197, 187)
    
    exp1_items = [
        "主导技能系统架构设计，支持50+种技能配置，实现技能连招、冷却、消耗等核心机制",
        "开发物理破坏系统，基于Chaos物理引擎实现建筑破坏效果，优化后同屏支持1000+物理物体",
        "负责单位皮肤系统，实现材质动态切换与LOD优化，渲染性能提升30%",
        "参与帧同步服务器开发，实现状态同步与回滚机制，网络延迟控制在50ms以内",
        "使用UE Profiler进行性能分析，通过Draw Call合并与材质优化，帧率从30fps提升至60fps",
    ]
    for item in exp1_items:
        p = right_cell.add_paragraph(style='List Bullet')
        p.space_after = Pt(4)
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(85, 85, 85)
    
    # 工作经历2
    p = right_cell.add_paragraph()
    p.space_before = Pt(15)
    p.space_after = Pt(3)
    run = p.add_run("2023.07 - 2024.12")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    p = right_cell.add_paragraph()
    p.space_after = Pt(8)
    run = p.add_run("U3D开发")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 51, 51)
    run = p.add_run("  |  ")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(153, 153, 153)
    run = p.add_run("宁波小伏科技有限公司")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(57, 197, 187)
    
    exp2_items = [
        "开发VR培训应用（Oculus Quest平台），实现手势识别与空间定位功能",
        "设计任务引导系统，采用行为树（Behavior Tree）实现NPC智能引导，用户完成率提升40%",
        "使用IK（反向动力学）技术实现肢体动画，支持15种工业操作动作",
        "优化VR渲染性能，通过注视点渲染（Foveated Rendering）技术，GPU负载降低25%",
    ]
    for item in exp2_items:
        p = right_cell.add_paragraph(style='List Bullet')
        p.space_after = Pt(4)
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(85, 85, 85)
    
    # 项目经验标题
    p = right_cell.add_paragraph()
    p.space_before = Pt(20)
    p.space_after = Pt(12)
    run = p.add_run("项目经验")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(57, 197, 187)
    
    # 添加下划线
    p = right_cell.add_paragraph()
    p.space_after = Pt(15)
    run = p.add_run("─" * 50)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(57, 197, 187)
    
    # 项目1
    p = right_cell.add_paragraph()
    p.space_after = Pt(5)
    run = p.add_run("战争警戒（War Alert）")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 51, 51)
    
    p = right_cell.add_paragraph()
    p.space_after = Pt(8)
    run = p.add_run("UE5 + C++ + 帧同步  |  团队规模：15人  |  开发周期：12个月")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(57, 197, 187)
    
    p = right_cell.add_paragraph()
    p.space_after = Pt(8)
    run = p.add_run("PVP RTS手游，支持实时多人对战。负责核心系统架构设计与性能优化：")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    project1_items = [
        "独立设计技能系统架构，支持技能组合、Buff/Debuff、伤害计算等复杂机制",
        "实现帧同步网络方案，解决延迟补偿与状态回滚问题，确保战斗一致性",
        "优化物理破坏效果，通过Chaos缓存与异步计算，同屏物理物体从200提升至1000+",
        "项目上线首月DAU 10万+，用户留存率45%",
    ]
    for item in project1_items:
        p = right_cell.add_paragraph(style='List Bullet')
        p.space_after = Pt(4)
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(85, 85, 85)
    
    # 项目2
    p = right_cell.add_paragraph()
    p.space_before = Pt(15)
    p.space_after = Pt(5)
    run = p.add_run("VR工业培训系统")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 51, 51)
    
    p = right_cell.add_paragraph()
    p.space_after = Pt(8)
    run = p.add_run("Unity + C# + Oculus SDK  |  团队规模：5人  |  开发周期：6个月")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(57, 197, 187)
    
    p = right_cell.add_paragraph()
    p.space_after = Pt(8)
    run = p.add_run("B端VR培训应用，用于制造业员工操作培训：")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    project2_items = [
        "开发任务引导系统，支持语音、文字、手势多种引导方式",
        "实现操作评分机制，记录用户操作轨迹并生成培训报告",
        "优化VR性能，确保90fps稳定帧率，降低晕动症发生率",
        "客户反馈培训效率提升60%，操作错误率降低80%",
    ]
    for item in project2_items:
        p = right_cell.add_paragraph(style='List Bullet')
        p.space_after = Pt(4)
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(85, 85, 85)
    
    # 项目3
    p = right_cell.add_paragraph()
    p.space_before = Pt(15)
    p.space_after = Pt(5)
    run = p.add_run("个人技术博客")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 51, 51)
    
    p = right_cell.add_paragraph()
    p.space_after = Pt(8)
    run = p.add_run("Hexo + Node.js + GitHub Actions")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(57, 197, 187)
    
    p = right_cell.add_paragraph()
    run = p.add_run("基于Hexo构建的静态博客系统，记录游戏开发技术心得。集成自动化部署、背景轮播、响应式设计等功能，使用GitHub Actions实现CI/CD自动化工作流。累计发布技术文章20+篇，涵盖Unity、UE5、性能优化等主题。")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    # 保存文档
    output_path = "resume_songhaoran.docx"
    doc.save(output_path)
    print(f"简历Word文档已生成：{output_path}")
    return output_path

if __name__ == "__main__":
    create_resume()
