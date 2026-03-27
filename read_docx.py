#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
读取docx文件内容
"""

from docx import Document

def read_docx():
    doc = Document('resume_songhaoran.docx')
    
    print("=== 文档段落内容 ===")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"{i}: {para.text}")
    
    print("\n=== 表格内容 ===")
    for i, table in enumerate(doc.tables):
        print(f"\n表格 {i}:")
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            print(row_text)

if __name__ == "__main__":
    read_docx()
