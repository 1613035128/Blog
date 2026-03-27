#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
完整读取docx文件内容
"""

from docx import Document

def read_docx():
    doc = Document('resume_songhaoran.docx')
    
    print("=== 完整文档内容 ===\n")
    
    # 读取所有段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"[{i}] {text}")
    
    print("\n=== 表格内容 ===")
    for table_idx, table in enumerate(doc.tables):
        print(f"\n--- 表格 {table_idx + 1} ---")
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            # 去重（合并单元格会重复）
            unique_cells = []
            for c in cells:
                if c and c not in unique_cells:
                    unique_cells.append(c)
            if unique_cells:
                print(f"  行{row_idx}: {' | '.join(unique_cells)}")

if __name__ == "__main__":
    read_docx()
